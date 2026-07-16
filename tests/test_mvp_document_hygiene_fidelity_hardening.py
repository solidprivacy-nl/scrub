from __future__ import annotations

from io import BytesIO
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document

from document_tools import anonymized_docx_from_original
from scrub_key import build_scrub_key
from scrub_key_document_reinsert import (
    DOCX_LIMITATIONS,
    reinsert_docx_bytes,
)


class UploadedBytes:
    def __init__(self, name: str, content: bytes):
        self.name = name
        self._content = bytes(content)

    def getvalue(self) -> bytes:
        return self._content


ROWS = [
    {
        "original_value": "BETROKKENE-TEST-A",
        "placeholder": "[PERSOON_1]",
        "entity_type": "PERSON",
        "type_label": "Naam",
        "source": "detected",
        "review_status": "auto_detected",
        "include": True,
        "timestamp": "2026-07-17T20:30:00Z",
    },
    {
        "original_value": "RECHTBANK TESTDAM",
        "placeholder": "[ORGANISATIE_01]",
        "entity_type": "ORGANIZATION",
        "type_label": "Organisatie",
        "source": "detected",
        "review_status": "auto_detected",
        "include": True,
        "timestamp": "2026-07-17T20:30:01Z",
    },
    {
        "original_value": "ZAAK-TEST-2026-001",
        "placeholder": "[ZAAKNUMMER_1]",
        "entity_type": "LEGAL_REFERENCE",
        "type_label": "Zaaknummer",
        "source": "manual",
        "review_status": "manual",
        "include": True,
        "timestamp": "2026-07-17T20:30:02Z",
    },
]

REPLACEMENTS = {
    row["original_value"]: row["placeholder"]
    for row in ROWS
}


def _scrub_key() -> dict:
    return build_scrub_key(ROWS, document_label="Synthetisch fidelitydossier")


def _source_docx() -> bytes:
    document = Document()
    document.add_paragraph("Body BETROKKENE-TEST-A")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Tabel ZAAK-TEST-2026-001"
    section = document.sections[0]
    section.header.paragraphs[0].text = "Header RECHTBANK TESTDAM"
    section.footer.paragraphs[0].text = "Footer ZAAK-TEST-2026-001"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _text_surfaces(docx_bytes: bytes) -> dict[str, str]:
    document = Document(BytesIO(docx_bytes))
    section = document.sections[0]
    return {
        "body": "\n".join(paragraph.text for paragraph in document.paragraphs),
        "table": "\n".join(
            paragraph.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
            for paragraph in cell.paragraphs
        ),
        "header": "\n".join(paragraph.text for paragraph in section.header.paragraphs),
        "footer": "\n".join(paragraph.text for paragraph in section.footer.paragraphs),
    }


def test_scrub_then_reinsert_restores_body_table_header_and_footer() -> None:
    source = _source_docx()
    scrubbed = anonymized_docx_from_original(
        UploadedBytes("synthetic.docx", source),
        REPLACEMENTS,
    )
    scrubbed_surfaces = _text_surfaces(scrubbed)

    assert "[PERSOON_1]" in scrubbed_surfaces["body"]
    assert "[ZAAKNUMMER_1]" in scrubbed_surfaces["table"]
    assert "[ORGANISATIE_01]" in scrubbed_surfaces["header"]
    assert "[ZAAKNUMMER_1]" in scrubbed_surfaces["footer"]

    result = reinsert_docx_bytes(scrubbed, _scrub_key())
    restored = _text_surfaces(result["docx_bytes"])

    assert restored["body"] == "Body BETROKKENE-TEST-A"
    assert restored["table"] == "Tabel ZAAK-TEST-2026-001"
    assert restored["header"] == "Header RECHTBANK TESTDAM"
    assert restored["footer"] == "Footer ZAAK-TEST-2026-001"
    assert result["replacement_count"] == 4
    assert result["placeholders_not_found"] == []
    assert result["unknown_placeholders"] == []
    assert result["validation_issues"] == []


def test_result_reports_processed_ooxml_parts() -> None:
    scrubbed = anonymized_docx_from_original(
        UploadedBytes("synthetic.docx", _source_docx()),
        REPLACEMENTS,
    )
    result = reinsert_docx_bytes(scrubbed, _scrub_key())

    assert "word/document.xml" in result["processed_parts"]
    assert any(part.startswith("word/header") for part in result["processed_parts"])
    assert any(part.startswith("word/footer") for part in result["processed_parts"])
    assert result["processed_part_count"] == len(result["processed_parts"])
    assert set(result["processed_parts"]) == set(result["part_texts"])
    assert result["local_only"] is True
    assert result["ai_processing"] is False
    assert result["cloud_processing"] is False


def test_comments_tracked_changes_metadata_and_split_nodes_remain_unsupported() -> None:
    limitation_text = " ".join(DOCX_LIMITATIONS).lower()

    assert "header/footer text are supported" in limitation_text
    assert "split" in limitation_text
    assert "comments" in limitation_text
    assert "tracked" in limitation_text
    assert "metadata" in limitation_text
    assert "not processed" in limitation_text


def test_unrelated_package_parts_are_preserved_byte_for_byte() -> None:
    source = _source_docx()
    marker_path = "customXml/synthetic-marker.xml"
    marker_bytes = b"<synthetic>KEEP-ME</synthetic>"
    enriched = BytesIO()
    with ZipFile(BytesIO(source), "r") as original, ZipFile(
        enriched,
        "w",
        ZIP_DEFLATED,
    ) as output:
        for entry in original.infolist():
            output.writestr(entry, original.read(entry.filename))
        output.writestr(marker_path, marker_bytes)

    scrubbed = anonymized_docx_from_original(
        UploadedBytes("synthetic.docx", enriched.getvalue()),
        REPLACEMENTS,
    )
    result = reinsert_docx_bytes(scrubbed, _scrub_key())

    with ZipFile(BytesIO(result["docx_bytes"]), "r") as restored:
        assert restored.read(marker_path) == marker_bytes


def test_invalid_header_xml_returns_validation_issue_without_partial_output() -> None:
    source = _source_docx()
    broken = BytesIO()
    with ZipFile(BytesIO(source), "r") as original, ZipFile(
        broken,
        "w",
        ZIP_DEFLATED,
    ) as output:
        for entry in original.infolist():
            data = original.read(entry.filename)
            if entry.filename.startswith("word/header") and entry.filename.endswith(".xml"):
                data = b"<w:hdr>"
            output.writestr(entry, data)

    result = reinsert_docx_bytes(broken.getvalue(), _scrub_key())

    assert result["reinserted"] is False
    assert result["replacement_count"] == 0
    assert result["validation_issues"]
    assert "header" in result["validation_issues"][0].lower()
    assert result["docx_bytes"] == broken.getvalue()


def test_pdf_ocr_and_restored_pdf_scope_are_not_added() -> None:
    source = Path("MVP_PHASE6_FALSE_NEGATIVE_GAP_TRIAGE.md").read_text(encoding="utf-8")

    assert "restored-TXT-only" in source
    assert "No OCR or restored-PDF implementation is authorized" in source
