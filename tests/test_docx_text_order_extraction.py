from io import BytesIO
from pathlib import Path

from docx import Document

from document_tools import extract_docx_text


def _save_docx_bytes(doc: Document) -> bytes:
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def test_docx_text_extraction_preserves_interleaved_paragraph_table_order():
    doc = Document()
    doc.add_paragraph("PAGINA_1_MARKER gewone tekst")

    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "PAGINA_2_MARKER tabeltekst"

    doc.add_paragraph("PAGINA_3_MARKER gewone tekst")

    roundtrip_doc = Document(BytesIO(_save_docx_bytes(doc)))
    text = extract_docx_text(roundtrip_doc)

    assert "PAGINA_1_MARKER" in text
    assert "PAGINA_2_MARKER" in text
    assert "PAGINA_3_MARKER" in text
    assert text.index("PAGINA_1_MARKER") < text.index("PAGINA_2_MARKER") < text.index("PAGINA_3_MARKER")


def test_docx_text_order_test_uses_synthetic_markers_only():
    doc = Document()
    doc.add_paragraph("PAGINA_1_MARKER gewone tekst")
    doc.add_paragraph("PAGINA_2_MARKER gewone tekst")
    extracted = extract_docx_text(Document(BytesIO(_save_docx_bytes(doc))))

    forbidden_real_markers = ["BSN", "IBAN", "Samir", "Marrakech", "dossiernummer"]
    for marker in forbidden_real_markers:
        assert marker not in extracted
