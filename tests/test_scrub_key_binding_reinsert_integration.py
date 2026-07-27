from __future__ import annotations

from copy import deepcopy
from io import BytesIO

from docx import Document

from scrub_key import build_scrub_key, scrub_key_to_json
from scrub_key_bound_export import build_bound_scrub_key
from scrub_key_document_reinsert import reinsert_docx_bytes, reinsert_txt_bytes
from scrub_key_import import build_scrub_key_import_result
from scrub_key_pdf_text_reinsert import reinsert_pdf_text_bytes
from scrub_key_reinsert import reinsert_from_scrub_key


BINDING_A = "BK7M4Q2XR5TD3W6YZ"
BINDING_B = "B234567ABCDEFGHJKL"
PERSON_A = f"[PERSOON_{BINDING_A}_01]"
DOSSIER_A = f"[DOSSIERNUMMER_{BINDING_A}_01]"
PERSON_B = f"[PERSOON_{BINDING_B}_01]"


def _rows(binding_id: str = BINDING_A) -> list[dict]:
    return [
        {
            "original_value": "BETROKKENE-TEST-A",
            "placeholder": f"[PERSOON_{binding_id}_01]",
            "entity_type": "PERSON",
            "type_label": "Naam",
            "source": "synthetic",
            "review_status": "synthetic_verified",
            "include": True,
            "timestamp": "2026-07-28T00:00:00Z",
        },
        {
            "original_value": "SYN-2026-0042",
            "placeholder": f"[DOSSIERNUMMER_{binding_id}_01]",
            "entity_type": "CASE_NUMBER",
            "type_label": "Dossiernummer",
            "source": "synthetic",
            "review_status": "synthetic_verified",
            "include": True,
            "timestamp": "2026-07-28T00:00:01Z",
        },
    ]


def _bound_key(binding_id: str = BINDING_A) -> dict:
    return build_bound_scrub_key(
        _rows(binding_id),
        document_binding_id=binding_id,
        document_label="Synthetische bindingtest",
    )


def _legacy_key() -> dict:
    rows = deepcopy(_rows())
    rows[0]["placeholder"] = "[PERSOON_01]"
    rows[1]["placeholder"] = "[DOSSIERNUMMER_01]"
    return build_scrub_key(rows, document_label="Synthetische legacytest")


def _docx_bytes(body: str, header: str = "", footer: str = "") -> bytes:
    document = Document()
    document.add_paragraph(body)
    section = document.sections[0]
    section.header.paragraphs[0].text = header
    section.footer.paragraphs[0].text = footer
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _docx_surfaces(content: bytes) -> dict[str, str]:
    document = Document(BytesIO(content))
    section = document.sections[0]
    return {
        "body": "\n".join(paragraph.text for paragraph in document.paragraphs),
        "header": "\n".join(paragraph.text for paragraph in section.header.paragraphs),
        "footer": "\n".join(paragraph.text for paragraph in section.footer.paragraphs),
    }


def test_bound_key_import_is_accepted_and_exposes_binding_metadata() -> None:
    result = build_scrub_key_import_result(scrub_key_to_json(_bound_key()))

    assert result["ok"] is True
    assert result["schema_version"] == "1.1"
    assert result["legacy_unbound"] is False
    assert result["key_binding_id"] == BINDING_A
    assert result["mapping_digest_valid"] is True
    assert result["item_count"] == 2


def test_bound_match_reinserts_and_reports_verified_document_match() -> None:
    result = reinsert_from_scrub_key(
        f"Cliënt {PERSON_A}; dossier {DOSSIER_A}.",
        _bound_key(),
    )

    assert result["text"] == "Cliënt BETROKKENE-TEST-A; dossier SYN-2026-0042."
    assert result["replacement_count"] == 2
    assert result["binding_status"] == "bound_match"
    assert result["replacement_allowed"] is True
    assert result["verified_document_match"] is True
    assert result["legacy_unbound"] is False
    assert result["mapping_digest_valid"] is True
    assert result["validation_issues"] == []


def test_wrong_bound_key_fails_closed_without_replacement() -> None:
    source = f"Cliënt {PERSON_A}."
    result = reinsert_from_scrub_key(source, _bound_key(BINDING_B))

    assert result["text"] == source
    assert result["replacement_count"] == 0
    assert result["binding_status"] == "binding_mismatch"
    assert result["replacement_allowed"] is False
    assert result["verified_document_match"] is False
    assert result["validation_issues"]


def test_mixed_document_bindings_fail_closed() -> None:
    source = f"Eerste {PERSON_A}; tweede {PERSON_B}."
    result = reinsert_from_scrub_key(source, _bound_key())

    assert result["text"] == source
    assert result["replacement_count"] == 0
    assert result["binding_status"] == "mixed_document_bindings"
    assert result["replacement_allowed"] is False
    assert result["document_binding_ids"] == [BINDING_B, BINDING_A]


def test_bound_key_for_legacy_document_fails_missing_document_binding() -> None:
    source = "Cliënt [PERSOON_01]."
    result = reinsert_from_scrub_key(source, _bound_key())

    assert result["text"] == source
    assert result["replacement_count"] == 0
    assert result["binding_status"] == "missing_document_binding"
    assert result["replacement_allowed"] is False


def test_legacy_key_for_bound_document_fails_closed() -> None:
    source = f"Cliënt {PERSON_A}."
    result = reinsert_from_scrub_key(source, _legacy_key())

    assert result["text"] == source
    assert result["replacement_count"] == 0
    assert result["binding_status"] == "legacy_key_for_bound_document"
    assert result["replacement_allowed"] is False
    assert result["legacy_unbound"] is True
    assert result["binding_warnings"]


def test_tampered_mapping_digest_fails_closed() -> None:
    key = _bound_key()
    key["mapping_digest"] = "0" * 64
    source = f"Cliënt {PERSON_A}."

    result = reinsert_from_scrub_key(source, key)

    assert result["text"] == source
    assert result["replacement_count"] == 0
    assert result["binding_status"] == "invalid_mapping_digest"
    assert result["replacement_allowed"] is False
    assert result["mapping_digest_valid"] is False
    assert result["validation_issues"]


def test_valid_legacy_key_remains_explicit_unverified_compatibility() -> None:
    result = reinsert_from_scrub_key(
        "Cliënt [PERSOON_01]; dossier [DOSSIERNUMMER_01].",
        _legacy_key(),
    )

    assert result["text"] == "Cliënt BETROKKENE-TEST-A; dossier SYN-2026-0042."
    assert result["replacement_count"] == 2
    assert result["binding_status"] == "legacy_unbound"
    assert result["replacement_allowed"] is True
    assert result["verified_document_match"] is False
    assert result["legacy_unbound"] is True
    assert result["binding_warnings"]


def test_txt_path_inherits_binding_gate() -> None:
    content = f"Cliënt {PERSON_A}.".encode("utf-8")
    matched = reinsert_txt_bytes(content, _bound_key())
    blocked = reinsert_txt_bytes(content, _bound_key(BINDING_B))

    assert matched["content_bytes"].decode("utf-8") == "Cliënt BETROKKENE-TEST-A."
    assert matched["binding_status"] == "bound_match"
    assert blocked["content_bytes"] == content
    assert blocked["replacement_count"] == 0
    assert blocked["binding_status"] == "binding_mismatch"


def test_docx_body_header_footer_restore_after_one_combined_binding_check() -> None:
    source = _docx_bytes(
        f"Body {PERSON_A}",
        header=f"Header {DOSSIER_A}",
        footer=f"Footer {PERSON_A}",
    )

    result = reinsert_docx_bytes(source, _bound_key())
    restored = _docx_surfaces(result["docx_bytes"])

    assert restored["body"] == "Body BETROKKENE-TEST-A"
    assert restored["header"] == "Header SYN-2026-0042"
    assert restored["footer"] == "Footer BETROKKENE-TEST-A"
    assert result["replacement_count"] == 3
    assert result["binding_status"] == "bound_match"
    assert result["verified_document_match"] is True


def test_docx_binding_mismatch_returns_original_bytes_without_partial_output() -> None:
    source = _docx_bytes(
        f"Body {PERSON_A}",
        header=f"Header {DOSSIER_A}",
        footer=f"Footer {PERSON_A}",
    )

    result = reinsert_docx_bytes(source, _bound_key(BINDING_B))

    assert result["docx_bytes"] == source
    assert result["replacement_count"] == 0
    assert result["binding_status"] == "binding_mismatch"
    assert result["replacement_allowed"] is False
    assert PERSON_A in result["text"]
    assert DOSSIER_A in result["text"]


def test_pdf_text_path_inherits_binding_gate(monkeypatch) -> None:
    import scrub_key_pdf_text_reinsert as pdf_module

    extracted = {
        "validation_issues": [],
        "unsupported_reason": None,
        "extracted_text": f"Cliënt {PERSON_A}.",
        "page_count": 1,
        "extraction_warnings": [],
    }
    monkeypatch.setattr(pdf_module, "extract_text_from_pdf_bytes", lambda _content: dict(extracted))

    matched = reinsert_pdf_text_bytes(b"%PDF-synthetic", _bound_key())
    blocked = reinsert_pdf_text_bytes(b"%PDF-synthetic", _bound_key(BINDING_B))

    assert matched["restored_text"] == "Cliënt BETROKKENE-TEST-A."
    assert matched["binding_status"] == "bound_match"
    assert blocked["restored_text"] == f"Cliënt {PERSON_A}."
    assert blocked["replacement_count"] == 0
    assert blocked["binding_status"] == "binding_mismatch"
    assert blocked["pdf_output"] is False
    assert blocked["ocr_used"] is False


def test_reinsert_does_not_mutate_bound_key() -> None:
    key = _bound_key()
    original = deepcopy(key)

    reinsert_from_scrub_key(f"Cliënt {PERSON_A}.", key)

    assert key == original
