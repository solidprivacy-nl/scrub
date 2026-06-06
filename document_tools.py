from io import BytesIO, StringIO
from html import escape
import csv
import re

try:
    from legal_reference_taxonomy import REFERENCE_ENTITY_TYPES
except Exception:
    REFERENCE_ENTITY_TYPES = []

import fitz  # PyMuPDF
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet


PLACEHOLDER_LABELS = {
    "PERSON": "PERSOON",
    "LOCATION": "LOCATIE",
    "ORGANIZATION": "ORGANISATIE",
    "EMAIL_ADDRESS": "EMAIL",
    "PHONE_NUMBER": "TELEFOON",
    "IBAN_CODE": "IBAN",
    "CREDIT_CARD": "BETAALKAART",
    "DATE_TIME": "DATUM",
    "URL": "URL",
    "IP_ADDRESS": "IP_ADRES",
    "GENERIC_PII": "VERTROUWELIJK",
    # Dutch/EU entities
    "NL_BSN": "BSN",
    "NL_POSTCODE": "POSTCODE",
    "NL_IBAN": "IBAN",
    "NL_KVK_NUMBER": "KVK_NUMMER",
    "NL_VAT_NUMBER": "BTW_NUMMER",
    "NL_PHONE_NUMBER": "TELEFOON",
    "NL_LICENSE_PLATE": "KENTEKEN",
    "NL_DRIVER_LICENSE": "RIJBEWIJS",
    "NL_BIG_NUMBER": "BIG_NUMMER",
    "NL_ADDRESS": "ADRES",
    "NL_DATE_OF_BIRTH": "GEBOORTEDATUM",
    # Dutch legal entities
    "NL_ECLI": "ECLI",
    "NL_LEGAL_CASE_NUMBER": "ZAAKNUMMER",
    "NL_ROLNUMMER": "ROLNUMMER",
    "NL_REKESTNUMMER": "REKESTNUMMER",
    "NL_PARKETNUMMER": "PARKETNUMMER",
    "NL_DOSSIER_NUMBER": "DOSSIERNUMMER",
    "NL_CLIENT_NUMBER": "CLIENTNUMMER",
    "NL_CJIB_NUMBER": "CJIB_NUMMER",
    "NL_POLICE_REPORT_NUMBER": "PV_NUMMER",
    "NL_INSURANCE_CLAIM_NUMBER": "SCHADE_OF_POLISNUMMER",
    "NL_INCIDENT_NUMBER": "INCIDENTNUMMER",
    "NL_CLAIM_NUMBER": "CLAIMNUMMER",
    "NL_OTHER_REFERENCE": "OVERIGE_REFERENTIE",
    "NL_LEGAL_PARTY_NAME": "PROCESPARTIJ",
    "NL_COURT_OR_AUTHORITY": "INSTANTIE",
    # Dutch legal reference taxonomy entities
    "NL_CLIENT_REFERENCE": "CLIENT_REFERENTIE",
    "NL_CASE_REFERENCE": "ZAAKREFERENTIE",
    "NL_INTERNAL_REFERENCE": "INTERNE_REFERENTIE",
    "NL_CONTEXTUAL_REFERENCE": "REFERENTIE",
    "NL_INVOICE_NUMBER": "FACTUURNUMMER",
    "NL_ORDER_OR_CONTRACT_NUMBER": "CONTRACT_OF_ORDERNUMMER",
    "NL_SCHOOL_REFERENCE": "SCHOOLREFERENTIE",
    "NL_CHILD_PROTECTION_REFERENCE": "JEUGD_OF_BESCHERMINGSREFERENTIE",
    "NL_EMPLOYMENT_REFERENCE": "ARBEIDSREFERENTIE",
    "NL_INSURANCE_REFERENCE": "VERZEKERINGSREFERENTIE",
    "NL_HEALTHCARE_REFERENCE": "ZORGREFERENTIE",
    "NL_POLICE_REFERENCE": "POLITIE_OF_OM_REFERENTIE",
    "NL_IMMIGRATION_REFERENCE": "VREEMDELINGENREFERENTIE",
    "NL_MUNICIPAL_REFERENCE": "BESTUURSREFERENTIE",
    "NL_REAL_ESTATE_REFERENCE": "VASTGOEDREFERENTIE",
    "NL_VEHICLE_REFERENCE": "VOERTUIG_OF_KENTEKENREFERENTIE",
    "NL_OBJECT_REFERENCE": "OBJECTREFERENTIE",
    "NL_SUSPICIOUS_REFERENCE_CANDIDATE": "MOGELIJKE_REFERENTIE",
    "NL_POSSIBLE_LICENSE_PLATE": "MOGELIJK_KENTEKEN",
}

STRUCTURED_ENTITY_TYPES = {
    "EMAIL_ADDRESS",
    "PHONE_NUMBER",
    "IBAN_CODE",
    "CREDIT_CARD",
    "DATE_TIME",
    "URL",
    "IP_ADDRESS",
    "NL_BSN",
    "NL_POSTCODE",
    "NL_IBAN",
    "NL_KVK_NUMBER",
    "NL_VAT_NUMBER",
    "NL_PHONE_NUMBER",
    "NL_LICENSE_PLATE",
    "NL_DRIVER_LICENSE",
    "NL_BIG_NUMBER",
    "NL_ADDRESS",
    "NL_DATE_OF_BIRTH",
    "NL_ECLI",
    "NL_LEGAL_CASE_NUMBER",
    "NL_ROLNUMMER",
    "NL_REKESTNUMMER",
    "NL_PARKETNUMMER",
    "NL_DOSSIER_NUMBER",
    "NL_CLIENT_NUMBER",
    "NL_CJIB_NUMBER",
    "NL_POLICE_REPORT_NUMBER",
    "NL_INSURANCE_CLAIM_NUMBER",
    "NL_LEGAL_PARTY_NAME",
    "NL_COURT_OR_AUTHORITY",
    *REFERENCE_ENTITY_TYPES,
}

# Common headings and document-structure words that should not become global replacements.
DOCUMENT_WORD_DENYLIST = {
    "chapter",
    "section",
    "article",
    "paragraph",
    "appendix",
    "annex",
    "schedule",
    "table",
    "figure",
    "introduction",
    "background",
    "summary",
    "conclusion",
    "scope",
    "purpose",
    "definitions",
    "agreement",
    "contract",
    "party",
    "parties",
    "client",
    "supplier",
    "service",
    "services",
    "project",
    "document",
    "version",
    "draft",
    "review",
    "confidential",
    "hoofdstuk",
    "paragraaf",
    "artikel",
    "bijlage",
    "inleiding",
    "samenvatting",
    "conclusie",
    "doel",
    "definities",
    "overeenkomst",
    "contract",
    "partij",
    "partijen",
    "klant",
    "leverancier",
    "dienst",
    "diensten",
    "project",
    "document",
    "versie",
    "concept",
    "vertrouwelijk",
    "rechtbank",
    "gerechtshof",
    "advocaat",
    "rechter",
    "griffier",
}


SENSITIVE_LEGAL_CONTEXT = {
    "strafzaak",
    "verdachte",
    "slachtoffer",
    "tenlastelegging",
    "veroordeling",
    "echtscheiding",
    "alimentatie",
    "omgangsregeling",
    "minderjarige",
    "jeugdzorg",
    "ontslag",
    "arbeidsongeschikt",
    "ziekte",
    "medisch",
    "asiel",
    "verblijfsvergunning",
    "faillissement",
    "curator",
}

LEGAL_ENTITY_TYPES = {
    "NL_ECLI",
    "NL_LEGAL_CASE_NUMBER",
    "NL_ROLNUMMER",
    "NL_REKESTNUMMER",
    "NL_PARKETNUMMER",
    "NL_DOSSIER_NUMBER",
    "NL_CLIENT_NUMBER",
    "NL_CJIB_NUMBER",
    "NL_POLICE_REPORT_NUMBER",
    "NL_INSURANCE_CLAIM_NUMBER",
    "NL_LEGAL_PARTY_NAME",
    "NL_COURT_OR_AUTHORITY",
    *REFERENCE_ENTITY_TYPES,
}


def uploaded_file_to_text(uploaded_file):
    """Convert uploaded .txt, .docx, or text-based .pdf to plain text."""
    filename = uploaded_file.name.lower()
    data = uploaded_file.getvalue()
    if filename.endswith(".txt"):
        return data.decode("utf-8", errors="ignore"), "txt"
    if filename.endswith(".docx"):
        doc = Document(BytesIO(data))
        return extract_docx_text(doc), "docx"
    if filename.endswith(".pdf"):
        return extract_pdf_text(data), "pdf"
    raise ValueError("Unsupported file type. Please upload .txt, .docx, or .pdf.")


def extract_docx_text(doc):
    parts = []
    for paragraph in iter_docx_paragraphs(doc):
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    return "\n".join(parts)


def extract_pdf_text(pdf_bytes):
    parts = []
    with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
        for page in doc:
            text = page.get_text("text", sort=True)
            if text.strip():
                parts.append(text)
    return "\n\n".join(parts)


def iter_docx_paragraphs(doc):
    """Yield paragraphs from body, tables, headers and footers."""
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        yield from iter_table_paragraphs(table)
    for section in doc.sections:
        header_footer_parts = [
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ]
        for part in header_footer_parts:
            for paragraph in part.paragraphs:
                yield paragraph
            for table in part.tables:
                yield from iter_table_paragraphs(table)


def iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested_table in cell.tables:
                yield from iter_table_paragraphs(nested_table)


def normalize_detected_text(value: str) -> str:
    return re.sub(r"\s+", " ", (value or "").strip())


def should_skip_detection(original: str, entity_type: str, score: float) -> bool:
    """Reduce false-positive global replacements in exported documents."""
    if not original:
        return True

    original_clean = normalize_detected_text(original)
    original_lower = original_clean.lower()

    if original_lower in DOCUMENT_WORD_DENYLIST:
        return True
    if not any(char.isalnum() for char in original_clean):
        return True
    if original_clean.isdigit() and entity_type not in STRUCTURED_ENTITY_TYPES:
        return True
    if len(original_clean) < 4 and entity_type not in STRUCTURED_ENTITY_TYPES:
        return True

    # Organization detection is useful, but often false-positives on headings.
    if entity_type == "ORGANIZATION":
        if score < 0.85:
            return True
        if len(original_clean.split()) == 1 and "." not in original_clean:
            return True

    if entity_type == "LOCATION" and score < 0.60:
        return True

    # Legal identifiers are usually safer to include once our recognizers find them.
    if entity_type in LEGAL_ENTITY_TYPES:
        return False

    return False


def placeholder_for_entity(entity_type: str, count: int) -> str:
    label = PLACEHOLDER_LABELS.get(entity_type, entity_type)
    return f"[{label}_{count:02d}]"


def build_placeholder_replacements(text, analyze_results):
    """Build stable placeholder suggestions from Presidio results."""
    counters = {}
    replacements = {}
    report_rows = []
    sorted_results = sorted(analyze_results, key=lambda r: (r.start, r.end))

    for result in sorted_results:
        original = normalize_detected_text(text[result.start : result.end])
        entity_type = result.entity_type
        score = float(getattr(result, "score", 0.0) or 0.0)

        if should_skip_detection(original, entity_type, score):
            continue
        if original in replacements:
            continue

        counters[entity_type] = counters.get(entity_type, 0) + 1
        placeholder = placeholder_for_entity(entity_type, counters[entity_type])
        replacements[original] = placeholder
        report_rows.append(
            {
                "entity_type": entity_type,
                "detected_text": original,
                "placeholder": placeholder,
                "score": round(score, 3),
            }
        )

    return replacements, report_rows


def apply_replacements_to_text(text, replacements):
    """Apply longest replacements first to avoid partial replacement problems."""
    output = text
    for original, placeholder in sorted(replacements.items(), key=lambda item: len(item[0]), reverse=True):
        output = output.replace(original, placeholder)
    return output


def anonymized_docx_from_original(uploaded_file, replacements):
    """Replace detected text inside the original .docx, preserving most formatting."""
    doc = Document(BytesIO(uploaded_file.getvalue()))
    for paragraph in iter_docx_paragraphs(doc):
        replace_in_paragraph_runs(paragraph, replacements)
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


def replace_in_paragraph_runs(paragraph, replacements):
    for original, placeholder in sorted(replacements.items(), key=lambda item: len(item[0]), reverse=True):
        while replace_once_in_runs(paragraph, original, placeholder):
            pass


def replace_once_in_runs(paragraph, old_text, new_text):
    """Replace one occurrence across Word runs."""
    if not paragraph.runs:
        return False
    full_text = "".join(run.text for run in paragraph.runs)
    start = full_text.find(old_text)
    if start == -1:
        return False
    end = start + len(old_text)
    current_pos = 0
    replacement_done = False

    for run in paragraph.runs:
        run_text = run.text
        run_start = current_pos
        run_end = current_pos + len(run_text)
        current_pos = run_end
        if run_end <= start or run_start >= end:
            continue
        local_start = max(start - run_start, 0)
        local_end = min(end - run_start, len(run_text))
        before = run_text[:local_start]
        after = run_text[local_end:]
        if not replacement_done:
            if end <= run_end:
                run.text = before + new_text + after
            else:
                run.text = before + new_text
            replacement_done = True
        else:
            if end <= run_end:
                run.text = after
            else:
                run.text = ""
    return True


def docx_from_text(text):
    doc = Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


def pdf_from_text(text):
    output = BytesIO()
    pdf = SimpleDocTemplate(output, pagesize=A4)
    styles = getSampleStyleSheet()
    normal = styles["Normal"]
    story = []
    for block in text.splitlines():
        safe_block = escape(block) if block.strip() else "&nbsp;"
        story.append(Paragraph(safe_block, normal))
        story.append(Spacer(1, 6))
    pdf.build(story)
    output.seek(0)
    return output.getvalue()


def replacement_report_csv(report_rows):
    output = StringIO()
    fieldnames = ["entity_type", "detected_text", "placeholder", "score", "source", "reason"]
    writer = csv.DictWriter(output, fieldnames=fieldnames, extrasaction="ignore")
    writer.writeheader()
    for row in report_rows:
        writer.writerow(row)
    return output.getvalue().encode("utf-8")


def scrub_report_txt(report_rows, profile: str, source_filename: str | None = None):
    """Build a simple local scrub report for legal review/admin evidence."""
    counts = {}
    legal_count = 0
    for row in report_rows:
        entity = str(row.get("entity_type", "UNKNOWN")) or "UNKNOWN"
        counts[entity] = counts.get(entity, 0) + 1
        if entity in LEGAL_ENTITY_TYPES:
            legal_count += 1

    lines = [
        "SolidPrivacy Scrub report",
        "==========================",
        "",
        f"Recognition profile: {profile}",
        f"Source file: {source_filename or 'text input / not specified'}",
        "Processing location: local app process / no external API call required by this recognizer pack",
        "Cloud/LLM use: none required for Dutch recognizer pack",
        "",
        "Detected entity counts:",
    ]

    if counts:
        for entity, count in sorted(counts.items()):
            lines.append(f"- {entity}: {count}")
    else:
        lines.append("- No entities included in the final replacement table.")

    lines.extend(
        [
            "",
            "Legal review note:",
            (
                "Manual review recommended: yes. Legal identifiers or personal identifiers may still be present "
                "if they were not detected or were unticked in the review table."
            ),
            "",
            "Scope warning:",
            "This is a scrubbing/pseudonymisation aid, not a guarantee of irreversible anonymisation.",
        ]
    )

    if legal_count:
        lines.insert(-4, f"Legal/matter identifiers included: {legal_count}")

    return "\n".join(lines).encode("utf-8")
