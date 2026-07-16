"""DOCX and PDF case runners for the MVP Phase 6 synthetic matrix."""

from __future__ import annotations

import re
from io import BytesIO
from typing import Any, Mapping

from docx import Document

from document_tools import (
    anonymized_docx_from_original,
    pdf_from_text,
    uploaded_file_to_text,
)
from docx_hygiene_audit import build_docx_hygiene_audit_report
from scrub_key_document_reinsert import reinsert_docx_bytes
from scrub_key_reinsert import detect_placeholders, reinsert_from_scrub_key

from mvp_phase6_workflow_core import (
    UploadedBytes,
    build_common_evidence,
)


def _normalise_whitespace(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def build_docx_case(case: Mapping[str, Any]) -> bytes:
    document = Document()
    for paragraph in case.get("body_paragraphs", []) or []:
        document.add_paragraph(str(paragraph))

    table_rows = case.get("table_rows", []) or []
    if table_rows:
        column_count = max(len(row) for row in table_rows)
        table = document.add_table(rows=len(table_rows), cols=column_count)
        for row_index, row in enumerate(table_rows):
            for column_index, value in enumerate(row):
                table.cell(row_index, column_index).text = str(value)

    section = document.sections[0]
    section.header.paragraphs[0].text = str(case.get("header_text") or "")
    section.footer.paragraphs[0].text = str(case.get("footer_text") or "")

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _expected_scope_is_supported(
    term: str,
    processed_parts: list[str],
    limitation_copy: str,
) -> bool:
    normalized = str(term).lower()
    if normalized in limitation_copy:
        return True
    if normalized in {"header", "headers"}:
        return any(part.startswith("word/header") for part in processed_parts)
    if normalized in {"footer", "footers"}:
        return any(part.startswith("word/footer") for part in processed_parts)
    return False


def run_docx_case(case: Mapping[str, Any]) -> dict[str, Any]:
    source_docx = build_docx_case(case)
    source_upload = UploadedBytes(f"{case['id']}.docx", source_docx)
    source_text, import_type = uploaded_file_to_text(source_upload)
    common = build_common_evidence(case, source_text)

    scrubbed_docx = anonymized_docx_from_original(
        source_upload,
        common["replacement_map"],
    )
    scrubbed_text, scrubbed_type = uploaded_file_to_text(
        UploadedBytes("scrubbed.docx", scrubbed_docx)
    )
    hygiene_report = build_docx_hygiene_audit_report(scrubbed_docx)
    reinsert_result = reinsert_docx_bytes(
        scrubbed_docx,
        common["scrub_key"],
    )
    restored_text, restored_type = uploaded_file_to_text(
        UploadedBytes("restored.docx", reinsert_result["docx_bytes"])
    )

    header_footer_text = [
        str(case.get("header_text") or ""),
        str(case.get("footer_text") or ""),
    ]
    formerly_expected_residuals = sorted(
        placeholder
        for original, placeholder in common["replacement_map"].items()
        if any(original in container for container in header_footer_text)
    )
    residuals = detect_placeholders(restored_text)
    resolved_header_footer_placeholders = sorted(
        placeholder
        for placeholder in formerly_expected_residuals
        if placeholder not in residuals
    )
    body_values = [
        original
        for original in common["replacement_map"]
        if not any(original in container for container in header_footer_text)
    ]
    body_values_present = all(value in restored_text for value in body_values)
    header_footer_values = [
        original
        for original in common["replacement_map"]
        if any(original in container for container in header_footer_text)
    ]
    header_footer_values_present = all(
        value in restored_text for value in header_footer_values
    )

    expected_findings = set(case.get("expected_hygiene_findings", []) or [])
    observed_findings = {
        finding.get("id")
        for finding in hygiene_report.get("findings", [])
        if finding.get("id")
    }
    audit_met = expected_findings.issubset(observed_findings)
    processed_parts = list(reinsert_result.get("processed_parts", []))
    limitation_copy = " ".join(reinsert_result.get("limitations", [])).lower()
    limitation_met = all(
        _expected_scope_is_supported(term, processed_parts, limitation_copy)
        for term in case.get("expected_reinsert_limitations", []) or []
    )
    roundtrip_complete = not residuals

    status = (
        "pass_with_known_limitations"
        if body_values_present
        and header_footer_values_present
        and audit_met
        and limitation_met
        and roundtrip_complete
        and not common["scrub_key_validation_issues"]
        else "fail"
    )

    return {
        "id": case["id"],
        "document_type": "docx",
        "status": status,
        "source_import_type": import_type,
        "scrubbed_import_type": scrubbed_type,
        "restored_import_type": restored_type,
        "source_text": source_text,
        "scrubbed_text": scrubbed_text,
        "restored_text": restored_text,
        "review_row_count": len(common["rows"]),
        "scrub_key_item_count": common["scrub_key"].get("item_count"),
        "scrub_key_validation_issues": common["scrub_key_validation_issues"],
        "reinsert_replacement_count": reinsert_result.get("replacement_count"),
        "body_roundtrip_values_present": body_values_present,
        "header_footer_roundtrip_values_present": header_footer_values_present,
        "expected_residual_placeholders": formerly_expected_residuals,
        "resolved_header_footer_placeholders": resolved_header_footer_placeholders,
        "residual_placeholders": residuals,
        "processed_parts": processed_parts,
        "hygiene_severity": hygiene_report.get("summary", {}).get("severity"),
        "hygiene_findings": sorted(observed_findings),
        "audit_expectations_met": audit_met,
        "known_limitations": list(reinsert_result.get("limitations", [])),
        "limitation_expectations_met": limitation_met,
        "roundtrip_complete": roundtrip_complete,
        "local_only": True,
        "ai_processing": False,
        "cloud_processing": False,
    }


def run_pdf_case(case: Mapping[str, Any]) -> dict[str, Any]:
    source_pdf = pdf_from_text(str(case.get("source_text") or ""))
    source_text, import_type = uploaded_file_to_text(
        UploadedBytes(f"{case['id']}.pdf", source_pdf)
    )
    common = build_common_evidence(case, source_text)

    scrubbed_pdf = pdf_from_text(common["scrubbed_text"])
    scrubbed_text, scrubbed_type = uploaded_file_to_text(
        UploadedBytes("scrubbed.pdf", scrubbed_pdf)
    )
    reinsert_result = reinsert_from_scrub_key(
        scrubbed_text,
        common["scrub_key"],
    )
    roundtrip_text_equal = (
        _normalise_whitespace(reinsert_result["text"])
        == _normalise_whitespace(source_text)
    )
    limitation_contract_met = (
        case.get("expected_reinsert_output_type") == "txt"
        and case.get("restored_pdf_supported") is False
        and case.get("ocr_supported") is False
    )
    status = (
        "pass_with_known_limitations"
        if roundtrip_text_equal
        and limitation_contract_met
        and not common["scrub_key_validation_issues"]
        else "fail"
    )

    return {
        "id": case["id"],
        "document_type": "pdf",
        "status": status,
        "source_import_type": import_type,
        "scrubbed_import_type": scrubbed_type,
        "source_text": source_text,
        "scrubbed_text": scrubbed_text,
        "restored_text": reinsert_result["text"],
        "review_row_count": len(common["rows"]),
        "scrub_key_item_count": common["scrub_key"].get("item_count"),
        "scrub_key_validation_issues": common["scrub_key_validation_issues"],
        "reinsert_replacement_count": reinsert_result.get("replacement_count"),
        "roundtrip_text_equal": roundtrip_text_equal,
        "reinsert_output_type": "txt",
        "restored_pdf_supported": False,
        "ocr_supported": False,
        "limitation_contract_met": limitation_contract_met,
        "known_limitations": [
            "PDF reinsert produces restored TXT only.",
            "No restored PDF output is supported.",
            "OCR and scanned/image-only PDFs are unsupported.",
        ],
        "local_only": True,
        "ai_processing": False,
        "cloud_processing": False,
    }
